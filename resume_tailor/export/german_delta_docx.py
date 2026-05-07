from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from resume_tailor.models import ResumeData
from resume_tailor.readers.docx_text import extract_first_image


ACCENT = RGBColor(49, 63, 74)
TEXT = RGBColor(36, 36, 36)
MUTED = RGBColor(96, 96, 96)


class GermanDeltaDocxRenderer:
    def __init__(self, template_docx: Path | None = None) -> None:
        self.template_docx = template_docx

    def render(self, resume: ResumeData, out_path: Path) -> None:
        doc = Document()
        section = doc.sections[0]
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.15)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

        _set_default_font(doc, "Aptos", 9)
        layout = doc.add_table(rows=1, cols=2)
        layout.autofit = False
        layout.columns[0].width = Cm(6.1)
        layout.columns[1].width = Cm(12.8)
        left = layout.cell(0, 0)
        right = layout.cell(0, 1)
        left.width = Cm(6.1)
        right.width = Cm(12.8)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _remove_table_borders(layout)

        photo = None
        if self.template_docx and self.template_docx.exists():
            photo = extract_first_image(self.template_docx, out_path.parent / "_template_photo.jpg")
        self._left_column(left, resume, photo)
        self._right_column(right, resume)

        doc.save(out_path)

    def _left_column(self, doc, resume: ResumeData, photo: Path | None) -> None:
        _clear_cell(doc)
        if photo:
            p = doc.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            p.add_run().add_picture(str(photo), width=Cm(4.55))
        else:
            _spacer(doc, 0.8)

        _name_block(doc, resume)
        _section_heading(doc, "KONTAKT")
        _body_line(doc, resume.phone)
        _body_line(doc, resume.email)
        _body_line(doc, resume.address)

        _section_heading(doc, "FÄHIGKEITEN")
        for skill in resume.skills[:10]:
            _body_line(doc, skill)

        _section_heading(doc, "SPRACHEN")
        for language in resume.languages:
            _body_line(doc, f"{language.name}: {language.level}")

    def _right_column(self, doc, resume: ResumeData) -> None:
        _clear_cell(doc)
        _section_heading(doc, "KURZPROFIL", top_space=0)
        _paragraph(doc, resume.profile, size=9, color=TEXT, after=7)

        _section_heading(doc, "BERUFSERFAHRUNG")
        for exp in resume.experiences:
            _role_line(doc, exp.title)
            _company_line(doc, exp.company)
            meta = exp.period
            if exp.location:
                meta += f", {exp.location}"
            _meta_line(doc, meta)
            for bullet in exp.bullets[:4]:
                _bullet(doc, bullet)
            _spacer(doc, 0.05)

        _section_heading(doc, "AUSBILDUNG")
        for edu in resume.education:
            _role_line(doc, f"{edu.degree}, {edu.institution}")
            detail = edu.period
            if edu.details:
                detail += f" | {edu.details}"
            _meta_line(doc, detail)


def _set_default_font(doc: Document, font_name: str, size: int) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size)
    style.font.color.rgb = TEXT


def _clear_cell(cell) -> None:
    for paragraph in list(cell.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)
    cell.add_paragraph()


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tag.set(qn("w:sz"), "0")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "FFFFFF")
        borders.append(tag)
    tbl_pr.append(borders)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
            if tc_borders is not None:
                tc_pr.remove(tc_borders)
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                tag = OxmlElement(f"w:{edge}")
                tag.set(qn("w:val"), "none")
                tag.set(qn("w:sz"), "0")
                tag.set(qn("w:color"), "FFFFFF")
                tc_borders.append(tag)
            tc_pr.append(tc_borders)


def _section_heading(cell, text: str, top_space: float = 0.22) -> None:
    if top_space:
        _spacer(cell, top_space)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    run.font.name = "Aptos Display"
    _bottom_border(p, ACCENT)


def _name_block(cell, resume: ResumeData) -> None:
    _spacer(cell, 0.18)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(resume.first_name)
    r1.font.size = Pt(15)
    r1.font.color.rgb = TEXT
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r2 = p.add_run(resume.last_name.upper())
    r2.bold = True
    r2.font.size = Pt(20)
    r2.font.color.rgb = ACCENT


def _role_line(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT


def _company_line(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT


def _meta_line(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def _body_line(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(8.6)
    run.font.color.rgb = TEXT


def _paragraph(cell, text: str, size: float, color: RGBColor, after: int = 4) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _bullet(cell, text: str) -> None:
    p = cell.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("• ")
    run.font.size = Pt(8.8)
    run.font.color.rgb = TEXT
    body = p.add_run(text)
    body.font.size = Pt(8.8)
    body.font.color.rgb = TEXT


def _spacer(cell, cm: float) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    run = p.add_run("")
    run.font.size = Pt(max(1, cm * 14))


def _bottom_border(paragraph, color: RGBColor) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), _color_hex(color))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _color_hex(color: RGBColor) -> str:
    return "".join(f"{channel:02X}" for channel in color)
